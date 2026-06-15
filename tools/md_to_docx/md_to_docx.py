#!/usr/bin/env python3
"""
Batch convert Markdown files to Word .docx documents.

Files:
  - md_paths.txt
  - output_dir.txt
  - processed_md_paths.txt

Dependencies:
  python3 -m pip install -r requirements.txt

Usage:
  python tools/md_to_docx/md_to_docx.py

Each non-empty line in md_paths.txt is treated as one Markdown file path.
Lines starting with # are ignored. The output directory is read from
output_dir.txt and is created automatically if needed.
"""

from __future__ import annotations

import os
import re
import sys
from pathlib import Path
from typing import Iterable

try:
    import markdown
    from bs4 import BeautifulSoup, NavigableString, Tag
    from docx import Document
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    from docx.shared import Inches, Pt, RGBColor
except ImportError as exc:  # pragma: no cover - handled at runtime
    print(f"Missing dependency: {exc}")
    print(
        "Install requirements with: "
        "python3 -m pip install python-docx markdown beautifulsoup4"
    )
    sys.exit(1)


URL_RE = re.compile(r"https?://[^\s<>\"]+")
BASE_DIR = Path(__file__).resolve().parent
REPO_ROOT = BASE_DIR.parent.parent
MD_PATHS_FILE = BASE_DIR / "md_paths.txt"
OUTPUT_DIR_FILE = BASE_DIR / "output_dir.txt"
PROCESSED_FILE = BASE_DIR / "processed_md_paths.txt"


def clean_user_path(raw: str) -> Path:
    value = raw.strip().strip('"').strip("'")
    path = Path(value).expanduser()
    if not path.is_absolute():
        path = (REPO_ROOT / path).resolve()
    else:
        path = path.resolve()
    return path


def read_list_file(path: Path) -> list[str]:
    if not path.exists():
        return []

    items: list[str] = []
    for line in path.read_text(encoding="utf-8").splitlines():
        stripped = line.strip()
        if stripped and not stripped.startswith("#"):
            items.append(stripped)
    return items


def sanitize_filename(name: str) -> str:
    name = re.sub(r'[\\/*?:"<>|]', "_", name)
    name = name.strip().strip(".")
    return name or "document"


def split_frontmatter(text: str) -> tuple[dict[str, str], str]:
    text = text.lstrip("\ufeff")
    if not text.startswith("---"):
        return {}, text

    lines = text.splitlines()
    if not lines or lines[0].strip() != "---":
        return {}, text

    end_idx = None
    for idx in range(1, len(lines)):
        if lines[idx].strip() == "---":
            end_idx = idx
            break

    if end_idx is None:
        return {}, text

    meta: dict[str, str] = {}
    for line in lines[1:end_idx]:
        stripped = line.strip()
        if not stripped or stripped.startswith("#"):
            continue
        if ":" not in line:
            continue
        key, value = line.split(":", 1)
        meta[key.strip().lower()] = value.strip().strip('"').strip("'")

    body = "\n".join(lines[end_idx + 1 :]).lstrip("\n")
    return meta, body


def is_url(value: str) -> bool:
    return value.startswith("http://") or value.startswith("https://")


def set_style_font(
    style,
    latin_font: str,
    east_asia_font: str,
    size: float | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
    color: RGBColor | None = None,
) -> None:
    style.font.name = latin_font
    if size is not None:
        style.font.size = Pt(size)
    if bold is not None:
        style.font.bold = bold
    if italic is not None:
        style.font.italic = italic
    if color is not None:
        style.font.color.rgb = color

    rpr = getattr(style._element, "rPr", None)
    if rpr is None:
        rpr = OxmlElement("w:rPr")
        style._element.append(rpr)
    rfonts = getattr(rpr, "rFonts", None)
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), latin_font)
    rfonts.set(qn("w:hAnsi"), latin_font)
    rfonts.set(qn("w:eastAsia"), east_asia_font)


def set_run_font(
    run,
    latin_font: str,
    east_asia_font: str,
    size: float | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
    color: RGBColor | None = None,
) -> None:
    run.font.name = latin_font
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if italic is not None:
        run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color

    rpr = getattr(run._element, "rPr", None)
    if rpr is None:
        rpr = OxmlElement("w:rPr")
        run._element.append(rpr)
    rfonts = getattr(rpr, "rFonts", None)
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), latin_font)
    rfonts.set(qn("w:hAnsi"), latin_font)
    rfonts.set(qn("w:eastAsia"), east_asia_font)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_horizontal_rule(paragraph, color: str = "A0A0A0") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def add_hyperlink(paragraph, text: str, url: str, *, bold=False, italic=False):
    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    r_pr.append(color)

    if bold:
        b = OxmlElement("w:b")
        r_pr.append(b)
    if italic:
        i = OxmlElement("w:i")
        r_pr.append(i)

    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    r_pr.append(u)

    run.append(r_pr)
    t = OxmlElement("w:t")
    t.text = text
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)
    return hyperlink


def choose_fonts() -> tuple[str, str, str]:
    if sys.platform == "darwin":
        return "Calibri", "PingFang SC", "Menlo"
    if os.name == "nt":
        return "Calibri", "Microsoft YaHei", "Consolas"
    return "Calibri", "Noto Sans CJK SC", "DejaVu Sans Mono"


def normalize_text(text: str) -> str:
    return text.replace("\r\n", "\n").replace("\r", "\n")


class MarkdownToDocx:
    def __init__(self, source_path: Path, output_dir: Path) -> None:
        self.source_path = source_path
        self.source_dir = source_path.parent
        self.output_dir = output_dir
        self.meta: dict[str, str] = {}
        self.body_font, self.east_asia_font, self.mono_font = choose_fonts()
        self.doc = Document()
        self.page_width = None
        self.usable_width = None

        self._setup_document()

    def _setup_document(self) -> None:
        section = self.doc.sections[0]
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.85)
        section.right_margin = Inches(0.85)
        section.header_distance = Inches(0.35)
        section.footer_distance = Inches(0.35)

        self.page_width = section.page_width
        # python-docx 的宽度运算结果会退化成 int，这里统一按 EMU 整数处理。
        self.usable_width = int(
            section.page_width - section.left_margin - section.right_margin
        )

        core = self.doc.core_properties
        core.title = self.source_path.stem
        core.author = "Codex"
        core.subject = "Markdown to Word conversion"

        styles = self.doc.styles

        normal = styles["Normal"]
        set_style_font(normal, self.body_font, self.east_asia_font, size=11)
        normal.paragraph_format.space_after = Pt(6)
        normal.paragraph_format.line_spacing = 1.25

        quote = self._ensure_style("Quote")
        set_style_font(
            quote,
            self.body_font,
            self.east_asia_font,
            size=10.5,
            italic=True,
            color=RGBColor(85, 85, 85),
        )
        quote.paragraph_format.left_indent = Inches(0.3)
        quote.paragraph_format.right_indent = Inches(0.1)
        quote.paragraph_format.space_before = Pt(2)
        quote.paragraph_format.space_after = Pt(6)

        code = self._ensure_style("Code Block")
        set_style_font(code, self.mono_font, self.east_asia_font, size=9.5)
        code.paragraph_format.left_indent = Inches(0.25)
        code.paragraph_format.right_indent = Inches(0.15)
        code.paragraph_format.space_before = Pt(4)
        code.paragraph_format.space_after = Pt(6)
        code.paragraph_format.keep_together = True

        table_text = self._ensure_style("Table Text")
        set_style_font(table_text, self.body_font, self.east_asia_font, size=9.5)
        table_text.paragraph_format.space_after = Pt(0)
        table_text.paragraph_format.line_spacing = 1.1

        for level in range(1, 7):
            style = styles[f"Heading {level}"]
            size = {1: 18, 2: 16, 3: 14, 4: 13, 5: 12, 6: 11}[level]
            set_style_font(
                style,
                self.body_font,
                self.east_asia_font,
                size=size,
                bold=True,
                color=RGBColor(31, 41, 55),
            )
            style.paragraph_format.space_before = Pt(8 if level == 1 else 6)
            style.paragraph_format.space_after = Pt(6 if level <= 3 else 4)
            style.paragraph_format.keep_with_next = True

    def _ensure_style(self, name: str):
        styles = self.doc.styles
        if name in styles:
            return styles[name]
        return styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)

    def load_markdown(self) -> str:
        raw = self.source_path.read_text(encoding="utf-8")
        self.meta, body = split_frontmatter(raw)
        if title := self.meta.get("title"):
            self.doc.core_properties.title = title
        if author := self.meta.get("author"):
            self.doc.core_properties.author = author
        return normalize_text(body)

    def convert(self, markdown_text: str) -> None:
        html = markdown.markdown(markdown_text, extensions=["extra"])
        soup = BeautifulSoup(html, "html.parser")
        self.render_container(soup.contents, list_level=0)

    def save(self) -> Path:
        self.output_dir.mkdir(parents=True, exist_ok=True)
        out_path = self._next_available_output_path()
        self.doc.save(str(out_path))
        return out_path

    def _next_available_output_path(self) -> Path:
        base_name = sanitize_filename(self.source_path.stem)
        candidate = self.output_dir / f"{base_name}.docx"
        index = 1
        while candidate.exists():
            candidate = self.output_dir / f"{base_name}-{index}.docx"
            index += 1
        return candidate

    def render_container(self, nodes: Iterable, list_level: int) -> None:
        for node in nodes:
            self.render_node(node, list_level=list_level)

    def render_node(self, node, list_level: int) -> None:
        if isinstance(node, NavigableString):
            text = str(node)
            if text.strip():
                paragraph = self.doc.add_paragraph(style="Normal")
                self.append_text(paragraph, text)
            return

        if not isinstance(node, Tag):
            return

        name = node.name.lower()

        if name in {"h1", "h2", "h3", "h4", "h5", "h6"}:
            level = int(name[1])
            paragraph = self.doc.add_paragraph(style=f"Heading {level}")
            self.render_inline_children(node.children, paragraph)
            return

        if name == "p":
            self.render_paragraph(node)
            return

        if name in {"ul", "ol"}:
            self.render_list(node, list_level)
            return

        if name == "blockquote":
            self.render_blockquote(node)
            return

        if name == "pre":
            self.render_code_block(node)
            return

        if name == "table":
            self.render_table(node)
            return

        if name == "hr":
            paragraph = self.doc.add_paragraph()
            add_horizontal_rule(paragraph)
            return

        if name == "img":
            paragraph = self.doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            self.add_image(paragraph, node)
            return

        if name in {"div", "section", "article", "main", "span", "body"}:
            self.render_container(node.children, list_level=list_level)
            return

        # Fallback: recurse into children to avoid dropping content.
        self.render_container(node.children, list_level=list_level)

    def render_paragraph(self, node) -> None:
        only_images = True
        for child in node.contents:
            if isinstance(child, NavigableString) and not str(child).strip():
                continue
            if isinstance(child, Tag) and child.name.lower() == "img":
                continue
            only_images = False
            break

        paragraph = self.doc.add_paragraph(style="Normal")
        if only_images:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.render_inline_children(node.children, paragraph)

    def render_blockquote(self, node) -> None:
        text = node.get_text("\n", strip=True)
        if not text:
            return
        paragraph = self.doc.add_paragraph(style="Quote")
        self.append_text(paragraph, text)

    def render_code_block(self, node) -> None:
        code = ""
        code_tag = node.find("code")
        if code_tag is not None:
            code = code_tag.get_text()
        else:
            code = node.get_text()
        code = code.rstrip("\n")
        if not code.strip():
            return

        paragraph = self.doc.add_paragraph(style="Code Block")
        run = paragraph.add_run(code)
        set_run_font(run, self.mono_font, self.east_asia_font, size=9.5)

        # Add light shading to make code blocks easier to scan.
        p_pr = paragraph._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "F6F8FA")
        p_pr.append(shd)

    def render_list(self, node, list_level: int) -> None:
        ordered = node.name.lower() == "ol"
        for li in node.find_all("li", recursive=False):
            head_nodes, tail_nodes = self.split_list_item(li)
            text = self.collect_plain_text(head_nodes).strip()

            if text:
                paragraph = self.doc.add_paragraph(style=self.list_style_name(ordered, list_level))
                if ordered:
                    # Word numbering styles handle numbering automatically.
                    pass
                self.render_inline_children(head_nodes, paragraph)
            else:
                paragraph = self.doc.add_paragraph(style=self.list_style_name(ordered, list_level))

            for child in tail_nodes:
                if isinstance(child, NavigableString):
                    if child.strip():
                        p = self.doc.add_paragraph(style="Normal")
                        p.paragraph_format.left_indent = Inches(0.25 * (list_level + 1))
                        self.append_text(p, str(child))
                    continue

                if not isinstance(child, Tag):
                    continue

                name = child.name.lower()
                if name in {"ul", "ol"}:
                    self.render_list(child, list_level + 1)
                elif name == "p":
                    p = self.doc.add_paragraph(style="Normal")
                    p.paragraph_format.left_indent = Inches(0.25 * (list_level + 1))
                    self.render_inline_children(child.children, p)
                elif name == "blockquote":
                    self.render_blockquote(child)
                elif name == "pre":
                    self.render_code_block(child)
                elif name == "table":
                    self.render_table(child)
                elif name == "img":
                    p = self.doc.add_paragraph()
                    p.paragraph_format.left_indent = Inches(0.25 * (list_level + 1))
                    self.add_image(p, child)
                else:
                    self.render_node(child, list_level=list_level + 1)

    def split_list_item(self, li) -> tuple[list, list]:
        head_nodes: list = []
        tail_nodes: list = []
        nested_started = False

        for child in li.children:
            if isinstance(child, NavigableString):
                if nested_started:
                    tail_nodes.append(child)
                else:
                    head_nodes.append(child)
                continue

            if not isinstance(child, Tag):
                continue

            name = child.name.lower()
            if name in {"ul", "ol"}:
                nested_started = True
                tail_nodes.append(child)
                continue

            if nested_started:
                tail_nodes.append(child)
            else:
                head_nodes.append(child)

        return head_nodes, tail_nodes

    def render_table(self, node) -> None:
        rows = self.extract_table_rows(node)
        if not rows:
            return

        col_count = max(len(row) for row in rows)
        if col_count == 0:
            return

        table = self.doc.add_table(rows=len(rows), cols=col_count)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False

        col_widths = self.calculate_table_column_widths(rows, col_count)
        for column, width in zip(table.columns, col_widths):
            column.width = width

        for row_idx, row in enumerate(rows):
            for col_idx in range(col_count):
                cell = table.cell(row_idx, col_idx)
                cell.width = col_widths[col_idx]
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
                cell.text = ""
                para = cell.paragraphs[0]
                para.style = "Table Text"

                if col_idx >= len(row):
                    continue

                cell_node = row[col_idx]
                self.render_inline_children(cell_node.children, para)
                if row_idx == 0 and cell_node.name.lower() in {"th", "td"}:
                    # Header-like row.
                    for run in para.runs:
                        run.bold = True
                    set_cell_shading(cell, "EDEDED")

        self.doc.add_paragraph()

    def calculate_table_column_widths(
        self, rows: list[list[Tag]], col_count: int
    ) -> list[int]:
        if col_count <= 0:
            return []

        weights = [8] * col_count
        for row in rows:
            for col_idx in range(min(col_count, len(row))):
                cell_node = row[col_idx]
                text = normalize_text(self.collect_plain_text(cell_node.children))
                if not text:
                    weight = 8
                else:
                    lines = [line.strip() for line in text.splitlines() if line.strip()]
                    if not lines:
                        weight = 8
                    else:
                        weight = max(len(line) for line in lines)
                        weight = min(max(weight, 8), 60)
                weights[col_idx] = max(weights[col_idx], weight)

        total_width = self.usable_width
        min_width = int(Inches(0.9))
        if col_count * min_width >= total_width:
            base = total_width // col_count
            widths = [base] * col_count
            for idx in range(total_width - sum(widths)):
                widths[idx % col_count] += 1
            return widths

        extra_width = total_width - (min_width * col_count)
        total_weight = sum(weights) or col_count
        widths = [
            min_width + (extra_width * weight // total_weight)
            for weight in weights
        ]

        remainder = total_width - sum(widths)
        if remainder > 0:
            for idx in sorted(range(col_count), key=lambda i: weights[i], reverse=True):
                if remainder <= 0:
                    break
                widths[idx] += 1
                remainder -= 1

        return widths

    def extract_table_rows(self, node) -> list[list[Tag]]:
        rows: list[list[Tag]] = []
        for child in node.children:
            if not isinstance(child, Tag):
                continue
            name = child.name.lower()
            if name == "tr":
                cells = [
                    grandchild
                    for grandchild in child.children
                    if isinstance(grandchild, Tag)
                    and grandchild.name.lower() in {"th", "td"}
                ]
                if cells:
                    rows.append(cells)
            elif name in {"thead", "tbody", "tfoot"}:
                for tr in child.find_all("tr", recursive=False):
                    cells = [
                        grandchild
                        for grandchild in tr.children
                        if isinstance(grandchild, Tag)
                        and grandchild.name.lower() in {"th", "td"}
                    ]
                    if cells:
                        rows.append(cells)
        return rows

    def render_inline_children(
        self, nodes: Iterable, paragraph, *, bold: bool = False, italic: bool = False
    ) -> None:
        for node in nodes:
            self.render_inline_node(node, paragraph, bold=bold, italic=italic)

    def render_inline_node(self, node, paragraph, *, bold=False, italic=False) -> None:
        if isinstance(node, NavigableString):
            self.append_text(paragraph, str(node), bold=bold, italic=italic)
            return

        if not isinstance(node, Tag):
            return

        name = node.name.lower()
        if name in {"strong", "b"}:
            self.render_inline_children(node.children, paragraph, bold=True, italic=italic)
            return
        if name in {"em", "i"}:
            self.render_inline_children(node.children, paragraph, bold=bold, italic=True)
            return
        if name == "code":
            text = node.get_text()
            run = paragraph.add_run(text)
            set_run_font(run, self.mono_font, self.east_asia_font, size=10)
            run.font.highlight_color = WD_COLOR_INDEX.GRAY_25
            return
        if name == "a":
            href = node.get("href", "").strip()
            label = node.get_text(" ", strip=True) or href
            if href:
                add_hyperlink(paragraph, label, href, bold=bold, italic=italic)
            else:
                self.append_text(paragraph, label, bold=bold, italic=italic)
            return
        if name == "br":
            paragraph.add_run().add_break()
            return
        if name == "img":
            self.add_image(paragraph, node)
            return
        if name in {"span", "u", "s", "del", "sup", "sub", "small", "mark", "kbd", "abbr", "cite"}:
            self.render_inline_children(node.children, paragraph, bold=bold, italic=italic)
            return
        if name in {"p", "div"}:
            self.render_inline_children(node.children, paragraph, bold=bold, italic=italic)
            paragraph.add_run().add_break()
            return

        self.render_inline_children(node.children, paragraph, bold=bold, italic=italic)

    def append_text(self, paragraph, text: str, *, bold=False, italic=False) -> None:
        if not text:
            return

        parts = []
        last = 0
        for match in URL_RE.finditer(text):
            if match.start() > last:
                parts.append(("text", text[last:match.start()]))
            parts.append(("url", match.group(0)))
            last = match.end()
        if last < len(text):
            parts.append(("text", text[last:]))

        if not parts:
            parts = [("text", text)]

        for kind, value in parts:
            if kind == "url":
                add_hyperlink(paragraph, value, value, bold=bold, italic=italic)
                continue

            segments = value.split("\n")
            for idx, segment in enumerate(segments):
                if segment:
                    run = paragraph.add_run(segment)
                    set_run_font(
                        run,
                        self.body_font,
                        self.east_asia_font,
                        bold=bold,
                        italic=italic,
                    )
                if idx < len(segments) - 1:
                    paragraph.add_run().add_break()

    def collect_plain_text(self, nodes: Iterable) -> str:
        chunks: list[str] = []
        for node in nodes:
            if isinstance(node, NavigableString):
                chunks.append(str(node))
            elif isinstance(node, Tag):
                name = node.name.lower()
                if name == "br":
                    chunks.append("\n")
                elif name == "img":
                    alt = node.get("alt", "").strip()
                    if alt:
                        chunks.append(alt)
                else:
                    chunks.append(self.collect_plain_text(node.children))
        return "".join(chunks)

    def list_style_name(self, ordered: bool, level: int) -> str:
        base = "List Number" if ordered else "List Bullet"
        if level <= 0:
            return base
        candidate = f"{base} {level + 1}"
        if candidate in self.doc.styles:
            return candidate
        return base

    def add_image(self, paragraph, node) -> None:
        src = node.get("src", "").strip()
        if not src:
            alt = node.get("alt", "").strip()
            if alt:
                self.append_text(paragraph, f"[Image: {alt}]")
            return

        if is_url(src):
            alt = node.get("alt", "").strip() or src
            self.append_text(paragraph, f"[External image omitted: {alt}]")
            return

        image_path = Path(src)
        if not image_path.is_absolute():
            image_path = (self.source_dir / image_path).resolve()

        if not image_path.exists():
            alt = node.get("alt", "").strip() or src
            self.append_text(paragraph, f"[Missing image: {alt}]")
            return

        try:
            shape = paragraph.add_run().add_picture(str(image_path))
            max_width = self.get_available_image_width(paragraph)
            if shape.width > max_width:
                scale = max_width / float(shape.width)
                shape.width = max_width
                shape.height = int(shape.height * scale)
        except Exception:
            alt = node.get("alt", "").strip() or src
            self.append_text(paragraph, f"[Image failed to load: {alt}]")

    def get_available_image_width(self, paragraph) -> int:
        parent = paragraph._p.getparent()
        if hasattr(parent, "width"):
            try:
                width = int(parent.width)
                if width > 0:
                    return width
            except Exception:
                pass

        available = self.usable_width
        left_indent = paragraph.paragraph_format.left_indent
        right_indent = paragraph.paragraph_format.right_indent
        if left_indent is not None:
            available -= int(left_indent)
        if right_indent is not None:
            available -= int(right_indent)

        return max(available, int(Inches(1)))


def main() -> int:
    if not MD_PATHS_FILE.exists():
        print(f"Missing path list: {MD_PATHS_FILE}")
        print("Create md_paths.txt with one Markdown path per line.")
        return 1

    if not OUTPUT_DIR_FILE.exists():
        print(f"Missing output config: {OUTPUT_DIR_FILE}")
        print("Create output_dir.txt with one output directory path.")
        return 1

    source_paths: list[Path] = []
    seen_sources: set[str] = set()
    for raw_path in read_list_file(MD_PATHS_FILE):
        source_path = clean_user_path(raw_path)
        source_key = str(source_path)
        if source_key in seen_sources:
            continue
        seen_sources.add(source_key)
        source_paths.append(source_path)

    if not source_paths:
        print(f"{MD_PATHS_FILE} contains no valid Markdown paths.")
        return 1

    output_dir_lines = read_list_file(OUTPUT_DIR_FILE)
    if not output_dir_lines:
        print(f"{OUTPUT_DIR_FILE} is empty.")
        print("Add one output directory path to output_dir.txt.")
        return 1
    output_dir = clean_user_path(output_dir_lines[0])

    processed_paths = load_processed_paths()
    pending_paths: list[Path] = []
    skipped_count = 0
    missing_count = 0

    for source_path in source_paths:
        source_key = str(source_path)
        if source_key in processed_paths:
            print(f"[自动跳过] 已转换过：{source_path}")
            skipped_count += 1
            continue
        if not source_path.exists() or not source_path.is_file():
            print(f"[跳过] 未找到 Markdown 文件：{source_path}")
            missing_count += 1
            continue
        if source_path.suffix.lower() != ".md":
            print(f"[跳过] 不是 .md 文件：{source_path}")
            missing_count += 1
            continue
        pending_paths.append(source_path)

    if not pending_paths:
        print("\n没有待转换的文件，程序退出")
        print(f"已跳过：{skipped_count} | 缺失/无效：{missing_count}")
        return 0

    output_dir.mkdir(parents=True, exist_ok=True)
    print(f"\n总计文件：{len(source_paths)} | 待转换：{len(pending_paths)}")
    print("-" * 60)

    converted_count = 0
    failed_count = 0

    for source_path in pending_paths:
        try:
            converter = MarkdownToDocx(source_path, output_dir)
            markdown_text = converter.load_markdown()
            converter.convert(markdown_text)
            out_path = converter.save()
            append_processed_path(source_path)
            processed_paths.add(str(source_path))
            converted_count += 1
            print(f"  ✅ {source_path} -> {out_path}")
        except Exception as exc:
            failed_count += 1
            print(f"  ❌ 处理失败：{source_path} | {str(exc)[:120]}")

    print("\n" + "=" * 60)
    print("全部任务执行完毕")
    print(f"已转换：{converted_count}")
    print(f"失败：{failed_count}")
    print(f"已跳过：{skipped_count}")
    print(f"缺失/无效：{missing_count}")
    print(f"输出目录：{output_dir}")
    print(f"已处理记录：{PROCESSED_FILE}")
    return 0


def load_processed_paths() -> set[str]:
    processed: set[str] = set()
    if not PROCESSED_FILE.exists():
        return processed
    for raw_path in read_list_file(PROCESSED_FILE):
        processed.add(str(clean_user_path(raw_path)))
    return processed


def append_processed_path(source_path: Path) -> None:
    PROCESSED_FILE.parent.mkdir(parents=True, exist_ok=True)
    with PROCESSED_FILE.open("a", encoding="utf-8") as f:
        f.write(str(source_path.resolve()) + "\n")


if __name__ == "__main__":
    raise SystemExit(main())
